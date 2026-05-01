from __future__ import annotations

import re
from pathlib import Path

from assignment_parser.models.base import Extractor
from assignment_parser.models.schema import Block, Document, Modality, SourceLocation
from assignment_parser.registry import register_extractor

try:
    import docx
except ImportError:  # pragma: no cover
    docx = None  # type: ignore[assignment]

_MD_HEADING_IN_PARA = re.compile(r"^\s*(#{1,6})\s+.+", re.DOTALL)


@register_extractor
class DocxExtractor(Extractor):
    modality = Modality.DOCX

    def can_handle(self, path: Path) -> bool:
        return path.suffix.lower() == ".docx"

    def extract(self, path: Path) -> Document:
        if docx is None:
            raise RuntimeError(
                "python-docx is required for .docx support. Install with: pip install python-docx"
            )
        document = docx.Document(str(path))
        blocks: list[Block] = []
        for p_idx, para in enumerate(document.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            loc = SourceLocation(paragraph_index=p_idx)
            style_name = (para.style.name if para.style else "") or ""
            heading_match = re.match(r"Heading\s+(\d+)", style_name, re.I)
            if heading_match:
                level = int(heading_match.group(1))
                blocks.append(
                    Block(
                        text=text,
                        location=loc,
                        kind="heading",
                        level=level,
                        metadata={"style": style_name},
                    )
                )
                continue
            m = _MD_HEADING_IN_PARA.match(text)
            if m:
                level = len(m.group(1))
                blocks.append(
                    Block(
                        text=text.strip(),
                        location=loc,
                        kind="heading",
                        level=level,
                        metadata={"style": "markdown_in_paragraph"},
                    )
                )
                continue
            blocks.append(
                Block(
                    text=text,
                    location=loc,
                    kind="text",
                    metadata={"style": style_name},
                )
            )
        return Document(
            blocks=blocks,
            modality=self.modality,
            source_path=str(path.resolve()),
            metadata={},
        )
