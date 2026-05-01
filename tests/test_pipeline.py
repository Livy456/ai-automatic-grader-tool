from __future__ import annotations

import json
from pathlib import Path

import pytest

from assignment_parser import parse, render
from assignment_parser.models.schema import Modality, TaskType
from assignment_parser.pipeline import parse as parse_fn


def test_markdown_imperative_bullets(tmp_path: Path) -> None:
    md = tmp_path / "hw.md"
    md.write_text(
        "Do the following:\n\n"
        "- remove duplicate rows from the table\n"
        "- plot the histogram of scores\n"
        "- click submit when finished\n",
        encoding="utf-8",
    )
    p = parse_fn(md)
    types = [t.task_type for t in p.tasks]
    assert TaskType.WRITTEN_RESPONSE in types
    assert len([t for t in p.tasks if "plot" in t.description.lower()]) >= 1
    assert not any("click" in t.description.lower() for t in p.tasks)


def test_notebook_todos_cell_index(tmp_path: Path) -> None:
    nb = {
        "nbformat": 4,
        "nbformat_minor": 5,
        "metadata": {},
        "cells": [
            {
                "cell_type": "code",
                "metadata": {},
                "source": ["# TODO: implement forward pass\nx = 1\n# FIXME: handle edge case\n"],
            }
        ],
    }
    path = tmp_path / "a.ipynb"
    path.write_text(json.dumps(nb), encoding="utf-8")
    p = parse_fn(path)
    assert len(p.tasks) == 1
    t = p.tasks[0]
    assert t.task_type == TaskType.CODE_TODO
    assert t.location.cell_index == 0
    assert "implement forward pass" in t.description
    assert "handle edge case" in t.description


def test_json_renderer_roundtrip(tmp_path: Path) -> None:
    md = tmp_path / "x.md"
    md.write_text("One paragraph only.\n", encoding="utf-8")
    p = parse_fn(md)
    raw = render(p, format="json")
    data = json.loads(raw)
    assert data["title"] == "x"
    assert data["modality"] == "markdown"
    assert "tasks" in data


def test_notebook_find_the_mistake(tmp_path: Path) -> None:
    nb = {
        "nbformat": 4,
        "nbformat_minor": 5,
        "metadata": {},
        "cells": [
            {
                "cell_type": "markdown",
                "metadata": {},
                "source": ["## Debug\n", "Find the mistake in the snippet below.\n"],
            },
            {"cell_type": "code", "metadata": {}, "source": ["print(1/0)\n"]},
        ],
    }
    path = tmp_path / "b.ipynb"
    path.write_text(json.dumps(nb), encoding="utf-8")
    p = parse_fn(path)
    assert any(t.task_type == TaskType.FIND_THE_MISTAKE for t in p.tasks)


def test_notebook_try_it_yourself_empty_code(tmp_path: Path) -> None:
    nb = {
        "nbformat": 4,
        "nbformat_minor": 5,
        "metadata": {},
        "cells": [
            {
                "cell_type": "markdown",
                "metadata": {},
                "source": ["## Exercise\n", "Try it yourself!\n"],
            },
            {"cell_type": "code", "metadata": {}, "source": []},
        ],
    }
    path = tmp_path / "c.ipynb"
    path.write_text(json.dumps(nb), encoding="utf-8")
    p = parse_fn(path)
    assert any(
        t.task_type == TaskType.CODE_TODO and "try it yourself" in t.description.lower()
        for t in p.tasks
    )


def test_markdown_likert_and_dedupe(tmp_path: Path) -> None:
    md = tmp_path / "likert.md"
    md.write_text(
        "This week we use a confidence scale.\n\n"
        "How confident are you in your solution? Rate from 1–5. ____\n\n"
        "How confident are you in your solution? Type your response below:\n",
        encoding="utf-8",
    )
    p = parse_fn(md)
    assert any(t.task_type == TaskType.SELF_RATING for t in p.tasks)
    assert not any(t.task_type == TaskType.WRITTEN_RESPONSE for t in p.tasks)


def test_markdown_critique_type_your_response(tmp_path: Path) -> None:
    md = tmp_path / "crit.md"
    md.write_text(
        "Please critique the visualization strengths and weaknesses in detail.\n\n"
        "Type your response here.\n",
        encoding="utf-8",
    )
    p = parse_fn(md)
    assert len(p.tasks) == 1
    assert p.tasks[0].task_type == TaskType.CRITIQUE


def test_markdown_renderer_summary_table(tmp_path: Path) -> None:
    md = tmp_path / "t.md"
    md.write_text("- plot the data\n", encoding="utf-8")
    p = parse_fn(md)
    out = render(p, format="markdown")
    assert "| Type | Description | Confidence | Section | Location |" in out


def test_unknown_extension_raises(tmp_path: Path) -> None:
    p = tmp_path / "x.unknown_ext"
    p.write_text("hi", encoding="utf-8")
    with pytest.raises(ValueError, match="No extractor"):
        parse_fn(p)


def test_missing_file_raises(tmp_path: Path) -> None:
    with pytest.raises(FileNotFoundError):
        parse_fn(tmp_path / "nope.ipynb")


def test_empty_notebook_no_tasks(tmp_path: Path) -> None:
    nb = {"nbformat": 4, "nbformat_minor": 5, "metadata": {}, "cells": []}
    path = tmp_path / "empty.ipynb"
    path.write_text(json.dumps(nb), encoding="utf-8")
    p = parse_fn(path)
    assert p.tasks == []


def test_no_headings_markdown_single_section(tmp_path: Path) -> None:
    md = tmp_path / "plain.md"
    md.write_text("Just one block of text.\n", encoding="utf-8")
    p = parse_fn(md)
    assert len(p.sections) == 1
    assert p.sections[0].title == "Document"


def test_transcript_srt(tmp_path: Path) -> None:
    from assignment_parser.segmenters.heading import walk_blocks_in_order

    srt = tmp_path / "t.srt"
    srt.write_text(
        "1\n"
        "00:00:00,000 --> 00:00:02,000\n"
        "Chapter: Introduction\n\n"
        "2\n"
        "00:00:02,000 --> 00:00:04,000\n"
        "Hello class.\n",
        encoding="utf-8",
    )
    p = parse_fn(srt)
    assert p.modality == Modality.VIDEO_TRANSCRIPT
    blocks = [b for _, b in walk_blocks_in_order(p.sections)]
    kinds = [b.kind for b in blocks]
    assert "heading" in kinds
    assert "text" in kinds


def test_docx_extractor(tmp_path: Path) -> None:
    import docx

    doc = docx.Document()
    doc.add_heading("Lab", level=1)
    doc.add_paragraph("### Custom heading in markdown style")
    doc.add_paragraph("Write your answer here about the experiment.")
    path = tmp_path / "lab.docx"
    doc.save(path)
    p = parse_fn(path)
    assert p.modality == Modality.DOCX
    assert any("Write your answer" in (t.raw_text or "") for t in p.tasks)


def test_unknown_renderer_raises(tmp_path: Path) -> None:
    md = tmp_path / "z.md"
    md.write_text("x\n", encoding="utf-8")
    p = parse_fn(md)
    with pytest.raises(ValueError, match="Unknown renderer"):
        render(p, format="not_a_format")


def test_parse_public_api_alias(tmp_path: Path) -> None:
    md = tmp_path / "api.md"
    md.write_text("- compute the mean\n", encoding="utf-8")
    p = parse(md)
    assert p.modality == Modality.MARKDOWN
